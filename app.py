import streamlit as st
import google.generativeai as genai
from groq import Groq
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from io import BytesIO
import PIL.Image
import PyPDF2
import re
import datetime

# --- 1. KONFIGURASI HALAMAN ---
st.set_page_config(
    page_title="Magis AI - Ignatian Pedagogy",
    page_icon="https://i.imgur.com/UUCgyfV.png",
    layout="wide",
    initial_sidebar_state="expanded",
)

# --- 2. CSS MODERN & DYNAMIC THEME ---
def inject_css():
    st.markdown("""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Plus+Jakarta+Sans:wght@400;500;600;700;800&display=swap');
    
    /* GLOBAL THEME */
    .stApp { 
        background-color: #F8F9FB !important; 
        color: #1E293B !important; 
        font-family: 'Plus Jakarta Sans', sans-serif; 
    }
    
    /* SIDEBAR STYLING */
    section[data-testid="stSidebar"] { 
        background-color: #FFFFFF !important; 
        border-right: 1px solid #E2E8F0;
        box-shadow: 4px 0 24px rgba(0,0,0,0.02);
    }
    
    /* TYPOGRAPHY UTAMA */
    h1, h2, h3 { color: #1B365D !important; letter-spacing: -0.5px; }
    
    /* HEADER JUDUL (GRADIENT TEXT) */
    .magis-title {
        font-weight: 800; 
        font-size: 48px; 
        background: linear-gradient(135deg, #1B365D 0%, #B8860B 100%);
        -webkit-background-clip: text; -webkit-text-fill-color: transparent;
        margin-bottom: 0px;
        line-height: 1.1;
        letter-spacing: -1px;
    }
    
    .magis-tagline {
        font-size: 18px;
        font-weight: 600;
        font-style: italic;
        color: #576F8E; 
        margin-bottom: 15px;
        border-left: 3px solid #DAA520;
        padding-left: 10px;
    }

    .magis-badge {
        display: inline-block;
        background-color: #E0F2FE;
        color: #0284C7;
        padding: 4px 12px;
        border-radius: 20px;
        font-size: 12px;
        font-weight: 700;
        margin-bottom: 20px;
    }
    
    /* CHAT BUBBLES (MODERN CARD STYLE) */
    .bubble-user {
        background: linear-gradient(135deg, #1B365D 0%, #2D4F85 100%);
        color: white; 
        padding: 20px; 
        border-radius: 20px 20px 4px 20px; 
        margin-left: auto; max-width: 85%;
        box-shadow: 0 10px 15px -3px rgba(27, 54, 93, 0.2);
        font-size: 15px;
        line-height: 1.6;
    }
    .bubble-ai {
        background-color: #FFFFFF; 
        color: #334155; 
        border: 1px solid #F1F5F9; 
        border-left: 5px solid #DAA520;
        padding: 24px; 
        border-radius: 4px 20px 20px 20px; 
        margin-right: auto; max-width: 95%;
        box-shadow: 0 4px 6px -1px rgba(0,0,0,0.05);
        font-size: 15px;
        line-height: 1.6;
    }
    
    /* INPUT AREA & FORM STYLING */
    .stTextArea textarea {
        background-color: #FFFFFF !important;
        border: 2px solid #E2E8F0 !important;
        border-radius: 12px !important;
        padding: 15px !important;
        font-size: 15px !important;
        transition: all 0.3s ease;
        box-shadow: inset 0 2px 4px rgba(0,0,0,0.02);
    }
    .stTextArea textarea:focus {
        border-color: #1B365D !important;
        box-shadow: 0 0 0 3px rgba(27, 54, 93, 0.1) !important;
    }
    
    /* BUTTON STYLING (GRADIENT & SHADOW) */
    div[data-testid="stForm"] button {
        background: linear-gradient(90deg, #1B365D 0%, #162B4A 100%);
        color: white;
        font-weight: 700;
        border-radius: 12px;
        padding: 10px 0;
        border: none;
        transition: transform 0.2s, box-shadow 0.2s;
        box-shadow: 0 4px 6px rgba(27, 54, 93, 0.2);
    }
    div[data-testid="stForm"] button:hover {
        transform: translateY(-2px);
        box-shadow: 0 10px 15px rgba(27, 54, 93, 0.3);
    }
    
    /* STATUS INDICATORS */
    .status-ok { color: #059669; font-weight: bold; font-size: 13px; border: 1px solid #059669; padding: 8px; border-radius: 8px; background: #ECFDF5; display: flex; align-items: center; gap: 5px;}
    .status-err { color: #DC2626; font-weight: bold; font-size: 13px; border: 1px solid #DC2626; padding: 8px; border-radius: 8px; background: #FEF2F2; display: flex; align-items: center; gap: 5px;}
    
    /* FOOTER STYLING */
    .sidebar-footer {
        text-align: center;
        margin-top: 30px;
        padding-top: 20px;
        border-top: 1px dashed #CBD5E1;
        color: #64748B;
        font-size: 12px;
        line-height: 1.5;
    }
    </style>
    """, unsafe_allow_html=True)

# --- 3. IGNATIAN DNA (OTAK UTAMA - VERSI 51.0 LENGKAP) ---
IGNATIAN_BASE_PROMPT = """
PERAN: 'Magis AI', asisten pedagogi dan pendamping spiritual khas Kolese Jesuit (Ignasian).

DNA SPIRITUAL & FILOSOFI (WAJIB DIINTEGRASIKAN):

1.  **IPP (Ignatian Pedagogical Paradigm):**
    - **Context:** Memahami dunia nyata siswa, keluarga, dan budaya.
    - **Experience:** Melibatkan rasa, akal budi, dan imajinasi (bukan hafalan semata).
    - **Reflection:** Menggali makna, nilai, dan kebenaran dari pengalaman.
    - **Action:** Dorongan untuk bertindak melayani sesama (Magis).
    - **Evaluation:** Menilai perkembangan karakter dan kompetensi secara utuh.

2.  **Cura Personalis:**
    - Perhatian personal pada keunikan tiap pribadi.
    - Mendengarkan dengan hati, memanusiakan, dan tidak menghakimi.

3.  **Universal Apostolic Preferences (UAP):**
    - Menunjukkan jalan menuju Tuhan (Latihan Rohani & Diskresi).
    - Berjalan bersama yang tersingkir (Solidaritas & Keadilan).
    - Menemani kaum muda menuju masa depan penuh harapan.
    - Merawat rumah kita bersama (Kesadaran Ekologis).

4.  **Ignatian Leadership & Values:**
    - **Self-Awareness:** Mengenal diri (kekuatan & kelemahan) dalam terang Ilahi.
    - **Ingenuity:** Cerdik dan inovatif dalam menghadapi tantangan zaman.
    - **Love & Heroism:** Kasih yang melayani lebih dari standar (Magis) dan keberanian memimpin.
    - **Profil Lulusan:** 4C (Competence, Conscience, Compassion, Commitment) + 1L (Leadership).

5.  **Diskresi (Pembedaan Roh) & Spiritualitas:**
    - **Pembedaan Roh:** Membantu membedakan gerak roh baik (Consolation: damai, sukacita, kasih) dan roh jahat (Desolation: gelisah, takut, putus asa).
    - **AMDG (Ad Maiorem Dei Gloriam):** Segala sesuatu demi kemuliaan Tuhan yang lebih besar.
    - **Nilai St. Petrus Canisius:** Ketekunan, kesetiaan pada Gereja, dan edukasi yang membebaskan.

6.  **Khazanah Kitab Suci & Musik Liturgi:**
    - **Alkitab Lengkap:** Akses penuh pada hikmat Perjanjian Lama & Baru. Gunakan referensi ayat secara akurat.
    - **Lagu Rohani:** Integrasikan lagu Katolik (Puji Syukur, Madah Bakti, Gregorian) dan lagu Kristiani yang relevan sebagai pengantar refleksi.

7.  **Kalender & Konteks Waktu (Nasional & Liturgi):**
    - **Kalender Nasional Indonesia:** Anda harus sadar akan hari libur nasional dan cuti bersama yang ditetapkan pemerintah Indonesia (misal: Lebaran, Natal, Waisak, Hari Kemerdekaan, dll) untuk membantu perencanaan konteks waktu pembelajaran.
    - **Kalender Liturgi Katolik Lengkap:** Anda memahami Tahun Liturgi (Tahun A/B/C, Tahun I/II), Masa (Adven, Natal, Prapaskah, Paskah, Biasa), Warna Liturgi, Bacaan Harian, dan Hari Raya/Pesta/Peringatan Wajib.

8.  **Tradisi, Sejarah, & Kekayaan Gereja:**
    - **Kumpulan Doa:** Menyediakan referensi doa-doa Katolik lengkap (Doa Dasar, Brevir, Novena, Jalan Salib, Rosario).
    - **Kisah Santo-Santa:** Memiliki pengetahuan mendalam tentang riwayat hidup, keutamaan, dan tanggal pesta para Santo-Santa pelindung (Hagiografi).
    - **Sejarah Alkitab:** Memahami konteks historis, geografis, dan budaya zaman Alkitab.
    - **Tata Perayaan Ekaristi (TPE):** Mengacu pada TPE Terkini (Revisi 2020/2021 di Indonesia) dalam menjelaskan struktur misa.

ATURAN OUTPUT:
- Gunakan bahasa Indonesia yang akademis, reflektif, namun hangat (Tone: Sahabat Pembimbing/Mentor).
- Setiap saran atau materi pembelajaran harus memiliki "kedalaman" (menyentuh aspek Conscience/Compassion), bukan hanya teknis.
- Matematika wajib menggunakan LaTeX ($...$).
- Jika membuat TABEL, gunakan format Markdown standard.
- Berikan output yang TERSTRUKTUR rapi.
- FOKUS pada teks dan konten materi. Jangan menyertakan tag gambar.
"""

# --- 4. ENGINE: AUTO-DISCOVERY & SELF HEALING ---
class AIProvider:
    def __init__(self, api_key):
        self.api_key = api_key
        self.provider_name = "None"
        self.client = None
        self.available_models = []
        self.active_model = None
        self.is_valid = False
        
        if not api_key: return

        if api_key.startswith("gsk_"):
            self.provider_name = "Groq"
            self._setup_groq()
        else:
            self.provider_name = "Google"
            self._setup_google()
        
    def _setup_groq(self):
        try:
            self.client = Groq(api_key=self.api_key)
            models = self.client.models.list()
            self.available_models = [m.id for m in models.data if 'llama' in m.id or 'mixtral' in m.id]
            self.available_models.sort(key=lambda x: '70b' in x, reverse=True)
            if self.available_models:
                self.active_model = self.available_models[0]
                self.is_valid = True
        except: pass

    def _setup_google(self):
        try:
            genai.configure(api_key=self.api_key)
            priorities = ['gemini-1.5-flash', 'gemini-1.5-pro', 'gemini-pro']
            self.available_models = priorities 
            try:
                all_models = genai.list_models()
                real_models = [m.name.replace("models/", "") for m in all_models if 'generateContent' in m.supported_generation_methods]
                if real_models: self.available_models = [p for p in priorities if p in real_models] + [m for m in real_models if m not in priorities]
            except: pass
            
            self.active_model = self.available_models[0] if self.available_models else 'gemini-1.5-flash'
            self.is_valid = True
        except: pass

    def generate_stream(self, history, prompt, system_config, image_input=None, lib_text=""):
        if not self.is_valid: yield "⚠️ Error: API Key bermasalah."; return

        full_system = f"{IGNATIAN_BASE_PROMPT}\n\n{system_config}"
        hist_str = "\n".join([f"{'USER' if m['role']=='user' else 'AI'}: {m['content']}" for m in history])
        
        # Tambahkan Tanggal Hari Ini agar AI sadar konteks waktu (untuk Liturgi/Libur)
        today_date = datetime.date.today().strftime("%A, %d %B %Y")
        
        final_prompt = f"TANGGAL HARI INI: {today_date}\n\nRIWAYAT CHAT:\n{hist_str}\n\nSUMBER PUSTAKA:\n{lib_text}\n\nPERMINTAAN USER:\n{prompt}"

        models_to_try = [self.active_model] + [m for m in self.available_models if m != self.active_model]
        success = False

        for model in models_to_try:
            if success: break
            try:
                if self.provider_name == "Google":
                    inputs = [f"SYSTEM_INSTRUCTION:\n{full_system}\n\nTASK:\n{final_prompt}"]
                    if image_input: inputs.append(image_input)
                    m = genai.GenerativeModel(model)
                    res = m.generate_content(inputs, stream=True)
                    for c in res: 
                        if c.text: yield c.text; success = True
                
                elif self.provider_name == "Groq":
                    if image_input: yield "ℹ️ [Groq: Gambar input diabaikan]\n"
                    stream = self.client.chat.completions.create(
                        messages=[{"role":"system","content":full_system},{"role":"user","content":final_prompt}],
                        model=model, stream=True
                    )
                    for c in stream:
                        txt = c.choices[0].delta.content
                        if txt: yield txt; success = True
            except: continue

# --- 5. LOGIC UI & HELPER (DOC ENGINE) ---
class DocEngine:
    @staticmethod
    def read(files):
        txt = ""; names = []
        for f in files:
            try:
                if f.name.endswith('.pdf'): txt += "".join([p.extract_text() for p in PyPDF2.PdfReader(f).pages])
                elif f.name.endswith('.docx'): txt += "\n".join([p.text for p in Document(f).paragraphs])
                elif f.name.endswith('.txt'): txt += f.getvalue().decode("utf-8")
                names.append(f.name)
            except: pass
        return txt, names

    @staticmethod
    def _set_table_borders(table):
        tbl = table._tbl
        for cell in tbl.iter_tcs():
            tcPr = cell.tcPr
            tcBorders = OxmlElement('w:tcBorders')
            top = OxmlElement('w:top'); top.set(qn('w:val'), 'single'); top.set(qn('w:sz'), '4')
            left = OxmlElement('w:left'); left.set(qn('w:val'), 'single'); left.set(qn('w:sz'), '4')
            bottom = OxmlElement('w:bottom'); bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '4')
            right = OxmlElement('w:right'); right.set(qn('w:val'), 'single'); right.set(qn('w:sz'), '4')
            tcBorders.append(top); tcBorders.append(left); tcBorders.append(bottom); tcBorders.append(right)
            tcPr.append(tcBorders)

    @staticmethod
    def _process_markdown_to_docx(doc, text):
        lines = text.split('\n')
        in_table = False
        table_data = []
        
        for line in lines:
            clean_line = line.strip()
            
            # Deteksi Tabel Markdown
            if "|" in clean_line and len(clean_line) > 2:
                if re.match(r'^\|?[\s-]+\|[\s-]+\|', clean_line): continue
                row_cells = [c.strip() for c in clean_line.split('|') if c.strip()]
                if row_cells:
                    if not in_table: in_table = True; table_data = [row_cells]
                    else: table_data.append(row_cells)
            else:
                if in_table:
                    if table_data:
                        table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                        table.style = 'Table Grid'
                        for r_idx, row_content in enumerate(table_data):
                            for c_idx, cell_content in enumerate(row_content):
                                if c_idx < len(table.columns):
                                    cell = table.cell(r_idx, c_idx)
                                    cell.text = cell_content.replace('**', '')
                        DocEngine._set_table_borders(table)
                        doc.add_paragraph()
                    in_table = False; table_data = []

                # Heading & Text Formatting
                if clean_line.startswith('### '): doc.add_heading(clean_line.replace('### ', ''), level=3)
                elif clean_line.startswith('## '): doc.add_heading(clean_line.replace('## ', ''), level=2)
                elif clean_line.startswith('# '): doc.add_heading(clean_line.replace('# ', ''), level=1)
                elif clean_line:
                    p = doc.add_paragraph()
                    parts = re.split(r'(\*\*.*?\*\*)', clean_line)
                    for part in parts:
                        if part.startswith('**') and part.endswith('**'):
                            run = p.add_run(part[2:-2]); run.bold = True
                        else: p.add_run(part.replace('$', ''))

    @staticmethod
    def create_word(history):
        doc = Document()
        title = doc.add_heading("Hasil Magis AI - Ignatian Pedagogy", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for msg in history:
            role_p = doc.add_heading(msg['role'].upper(), level=2)
            role_p.runs[0].font.color.rgb = RGBColor(27, 54, 93)
            DocEngine._process_markdown_to_docx(doc, msg['content'])
            doc.add_paragraph("-" * 20)

        bio = BytesIO(); doc.save(bio); return bio

# --- 6. APLIKASI UTAMA ---
if 'history' not in st.session_state: st.session_state.history = []
if 'library' not in st.session_state: st.session_state.library = {"text":"", "files":[]}

inject_css()

# API Handling
api_key = None
try:
    if "GOOGLE_API_KEY" in st.secrets: api_key = st.secrets["GOOGLE_API_KEY"]
    elif "GROQ_API_KEY" in st.secrets: api_key = st.secrets["GROQ_API_KEY"]
except: pass 

# --- SIDEBAR & SMART INPUT LOGIC ---
with st.sidebar:
    st.markdown("""
        <div style="text-align: center; margin-bottom: 20px;">
            <img src="https://i.imgur.com/UUCgyfV.png" width="110" style="filter: drop-shadow(0px 4px 6px rgba(0,0,0,0.1));">
        </div>
    """, unsafe_allow_html=True)
    
    if not api_key: 
        st.info("🔐 Masukkan Kunci Akses")
        api_key = st.text_input("API Key", type="password", label_visibility="collapsed")
        
    provider = AIProvider(api_key)
    if provider.is_valid: st.markdown(f"<div class='status-ok'>✅ Sistem {provider.provider_name} Terhubung</div>", unsafe_allow_html=True)
    else: st.markdown("<div class='status-err'>⚠️ Menunggu API Key</div>", unsafe_allow_html=True)
    
    st.markdown("---")
    
    # MENU DENGAN IKON (UPDATE: Mode Bebas)
    mode = st.selectbox("📌 Pilih Divisi Pelayanan", 
                        ["Akademik (Pedagogi)", "Pastoral & Diskresi", "Manajemen Sekolah", "✨ Obrolan Bebas (General Chat)"])
    config_details = ""
    auto_prompt_template = "" 
    
    if mode == "Akademik (Pedagogi)":
        st.markdown("#### 🎓 Konfigurasi Akademik")
        with st.expander("📚 Kelas & Materi", expanded=True):
            input_kelas = st.selectbox("Jenjang Kelas", ["7 SMP", "8 SMP", "9 SMP", "10 SMA (Fase E)", "11 SMA (Fase F)", "12 SMA (Fase F)"])
            input_mapel = st.text_input("Mata Pelajaran", placeholder="Misal: Sejarah Indonesia")
            input_kd = st.text_area("Kompetensi Dasar (KD) / CP", placeholder="Paste CP/Tujuan Pembelajaran...", height=80)
            
        with st.expander("🧠 Parameter Soal & Tugas"):
            input_bloom = st.multiselect("Level Kognitif (Bloom)", 
                                         ["C1 (Mengingat)", "C2 (Memahami)", "C3 (Menerapkan)", "C4 (Menganalisis)", "C5 (Mengevaluasi)", "C6 (Mencipta)"],
                                         default=["C4 (Menganalisis)", "C5 (Mengevaluasi)"])
            input_difficulty = st.select_slider("Tingkat Kesulitan", options=["Mudah", "Sedang", "HOTS (Sulit)", "Olimpiade"])
            
        with st.expander("🎨 Gaya & Pendekatan Ignasian"):
            input_gaya = st.selectbox("Gaya Bahasa", ["Formal Akademis", "Sokratik (Bertanya Balik)", "Storytelling (Naratif)", "Simpel & Lugas"])
            input_ipp_focus = st.multiselect("Fokus IPP", ["Context", "Experience", "Reflection", "Action", "Evaluation"], default=["Reflection"])

        config_details = f"KONFIGURASI: Kelas {input_kelas}, Mapel {input_mapel}, Gaya {input_gaya}, IPP {','.join(input_ipp_focus)}"
        
        auto_prompt_template = (
            f"Saya guru {input_mapel} untuk kelas {input_kelas}. \n"
            f"Topik: {input_kd if input_kd else '[Isi Topik]'}. \n\n"
            f"Tolong buatkan [Rencana Pembelajaran / 5 Soal PG / Studi Kasus] "
            f"dengan level kognitif {', '.join(input_bloom)} dan tingkat kesulitan {input_difficulty}. "
            f"Tekankan pada aspek {', '.join(input_ipp_focus)}."
        )

    elif mode == "Pastoral & Diskresi":
        st.markdown("#### 🕊️ Pendampingan Pastoral")
        with st.expander("❤️ Konteks Konseling", expanded=True):
            pas_subjek = st.selectbox("Subjek", ["Siswa", "Guru/Karyawan", "Orang Tua", "Alumni"])
            
            opsi_masalah = ["Akademik", "Keluarga", "Pencarian Jati Diri", "Keputusan Besar (Diskresi)", "Kejenuhan/Burnout", "Lainnya (Tulis Sendiri)..."]
            pilihan_masalah = st.selectbox("Isu Utama", opsi_masalah)
            
            if pilihan_masalah == "Lainnya (Tulis Sendiri)...":
                pas_masalah = st.text_input("Tuliskan Isu Spesifik:", placeholder="Misal: Konflik dengan teman sebaya...")
            else:
                pas_masalah = pilihan_masalah
            
            pas_metode = st.radio("Metode Pendampingan", ["Mendengarkan (Listening)", "Diskresi (Pembedaan Roh)", "Examen (Refleksi Harian)"])
        
        config_details = f"KONFIGURASI PASTORAL: Subjek {pas_subjek}, Masalah {pas_masalah}, Metode {pas_metode}"
        
        auto_prompt_template = (
            f"Saya sedang mendampingi seorang {pas_subjek} yang mengalami pergumulan tentang {pas_masalah if pas_masalah else '[Isi Masalah]'}. \n\n"
            f"Mohon berikan panduan percakapan atau refleksi menggunakan pendekatan {pas_metode}. "
            f"Tujuannya adalah membantu subjek menemukan kedamaian (konsolasi) dan mengambil keputusan yang tepat."
        )
        
    elif mode == "Manajemen Sekolah":
        st.markdown("#### 💼 Manajemen Sekolah")
        man_jenis = st.selectbox("Jenis Dokumen", ["Surat Resmi", "Proposal Kegiatan", "Pidato/Sambutan", "Email Internal"])
        man_tone = st.select_slider("Nada Bicara", options=["Tegas & Formal", "Persuasif", "Apresiatif", "Instruktif"])
        man_topik = st.text_input("Topik/Acara", placeholder="Misal: Hari Guru")
        
        config_details = f"KONFIGURASI MANAJEMEN: Dokumen {man_jenis}, Tone {man_tone}"
        
        auto_prompt_template = (
            f"Buatkan draf {man_jenis} bertema '{man_topik if man_topik else '[Isi Topik]'}'. \n\n"
            f"Gunakan nada bicara yang {man_tone}. "
            f"Pastikan struktur dokumen rapi dan sesuai standar institusi pendidikan Jesuit."
        )
    
    else: # MODE OBROLAN BEBAS
        st.markdown("#### 💬 Diskusi Terbuka")
        st.info("Mode ini membebaskan Anda berdiskusi topik apapun dengan perspektif Ignasian.")
        
        config_details = "KONFIGURASI: Mode Diskusi Bebas. Berperanlah sebagai 'Ignatian Friend' yang bijaksana, mendalam, dan suportif."
        auto_prompt_template = "" 

    # FOOTER AREA
    st.markdown("<br>", unsafe_allow_html=True)
    if st.button("🔄 Reset Sesi Chat", use_container_width=True): 
        st.session_state.history = []
        st.rerun()
        
    st.markdown("""
        <div class="sidebar-footer">
            <strong>Magis AI v51.0</strong><br>
            Design by: Albertus Henny Setyawan<br>
            Kolese Kanisius Jakarta | 2026
        </div>
    """, unsafe_allow_html=True)

# --- MAIN UI ---
c1,c2 = st.columns([3,1])
with c1: 
    st.markdown(f'''
    <div class="magis-title">MAGIS AI</div>
    <div class="magis-tagline">Mitra Diskresi Guru Ignasian</div>
    <div class="magis-badge">Mode Aktif: {mode}</div>
    ''', unsafe_allow_html=True)

with st.expander("📂 Upload Dokumen & Materi Referensi", expanded=False):
    st.markdown("Upload RPP, E-Book, atau Gambar Soal untuk dianalisis AI.")
    files = st.file_uploader("Pilih file (PDF, Docx, TXT)", accept_multiple_files=True)
    img_up = st.file_uploader("Upload Gambar (Jika perlu)", type=['png','jpg'])
    if files:
        t, n = DocEngine.read(files)
        st.session_state.library = {"text": t, "files": n}
        st.success(f"📚 {len(n)} dokumen berhasil dipelajari.")

# --- CHAT DISPLAY ---
chat_container = st.container()
with chat_container:
    for m in st.session_state.history:
        st.markdown(f"<div class='{'bubble-user' if m['role']=='user' else 'bubble-ai'}'>{m['content'].replace('[DOC_CONTEXT]','')}</div>", unsafe_allow_html=True)
    
    # Spacer kosong
    st.markdown("<div style='height: 50px;'></div>", unsafe_allow_html=True)

# --- SMART INPUT AREA (LEBAR & NYAMAN) ---
st.markdown("---")
st.markdown("### ✍️ Area Kerja")

with st.form(key='smart_input_form', clear_on_submit=True):
    # KEY TRICK: Hash template agar auto-refresh
    prompt_key = f"input_{hash(auto_prompt_template)}" 
    
    user_in = st.text_area(
        "Tulis instruksi atau pesan Anda:", 
        value=auto_prompt_template, 
        height=250, 
        key=prompt_key,
        placeholder="Ketik pesan Anda di sini..." # Placeholder muncul jika template kosong
    )
    
    col_act1, col_act2 = st.columns([1, 5])
    with col_act1:
        submitted = st.form_submit_button("🚀 KIRIM PERINTAH", use_container_width=True)
    with col_act2:
        if mode != "✨ Obrolan Bebas (General Chat)":
            st.caption("💡 *Tip: Edit draf di atas sesuai kebutuhan spesifik Anda.*")
        else:
            st.caption("💡 *Tip: Silakan berdiskusi bebas, Magis AI siap menjadi teman berpikir.*")

# LOGIC PEMROSESAN
if submitted and user_in and provider.is_valid:
    final_msg = user_in
    if st.session_state.library["text"]: final_msg += " [DOC_CONTEXT]"
    st.session_state.history.append({"role":"user", "content":final_msg})
    st.rerun()

if st.session_state.history and st.session_state.history[-1]['role'] == 'user':
    with st.spinner("✨ Sedang meracik materi dengan perspektif Ignasian..."):
        full_res = ""
        box = st.empty()
        last_usr = st.session_state.history[-1]['content']
        img_data = PIL.Image.open(img_up) if img_up else None
        
        for chk in provider.generate_stream(st.session_state.history[:-1], last_usr, config_details, img_data, st.session_state.library["text"]):
            full_res += chk
            box.markdown(f"<div class='bubble-ai'>{full_res}</div>", unsafe_allow_html=True)
            
        st.session_state.history.append({"role":"assistant", "content":full_res})
        st.rerun()

if st.session_state.history:
    st.markdown("### 📥 Ekspor Hasil")
    docx = DocEngine.create_word(st.session_state.history)
    st.download_button(
        label="Download Dokumen Word (.docx)", 
        data=docx, 
        file_name="Hasil-MagisAI.docx", 
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True
    )